import sys
import os
import re
import comtypes.client
import logging
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import difflib

#logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def doc_to_docx(doc_path):
    try:
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(doc_path)
        doc.Activate()

        new_file_abs, _ = os.path.splitext(os.path.abspath(doc_path))
        new_file_abs += '.docx'

        word.ActiveDocument.SaveAs(new_file_abs, FileFormat=16)
        doc.Close()
        word.Quit()

        if os.path.exists(doc_path):
            os.remove(doc_path)

        return new_file_abs
    except Exception as e:
        logging.error(f"Error converting {doc_path} to .docx: {e}")
        return None

class WordProcessor:
    def __init__(self):
        self.template_text = None
        self.documents = {}  # 字典来存储姓名(文件名)-内容-评分

    def set_template(self, template_text):
        if template_text:
            self.template_text = template_text
            logging.info("Template text has been set.")
        else:
            logging.warning("Attempted to set empty template text.")

    def read_docx(self, file_path):
        """
        读取DOCX文件中的所有文本内容，包括段落和表格中的文本。
        参数:file_path (str): 要读取的DOCX文件的路径。
        返回:str: DOCX文件的全部文本内容，如果发生错误则返回None。
        """
        try:
            doc = Document(file_path)
            full_text = []

            # 读取段落中的文本
            for para in doc.paragraphs:
                full_text.append(para.text)

            # 读取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            # 将所有文本合并成一个字符串
            document_text = '\n'.join(full_text)
            logging.info(f"Successfully read document: {file_path}")
            return document_text
        except Exception as e:
            # 记录错误详情
            logging.error(f"Error reading {file_path}: {e}")
            return None

    def remove_template_content(self, student_text, template_text):
        """
        从学生文档中移除模板内容。
        使用difflib库来查找和移除最长的匹配字符串，这种方法可以处理一些格式上的小差异。
        """
        s = difflib.SequenceMatcher(None, student_text, template_text)
        match = s.find_longest_match(0, len(student_text), 0, len(template_text))

        # 如果找到了足够长的匹配内容，则移除它
        if match.size > 20:  # 可以调整这个阈值以避免移除过短的匹配，可能是误匹配
            student_text = student_text[:match.a] + student_text[match.a + match.size:]
        return student_text

    def preprocess_text(self, text):
        """Preprocess text by removing template content."""
        if self.template_text and text:
            return self.remove_template_content(text, self.template_text)
        return text

    def doc_to_docx(self, doc_path):
        try:
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(doc_path)
            doc.Activate()

            new_file_abs, _ = os.path.splitext(os.path.abspath(doc_path))
            new_file_abs += '.docx'

            word.ActiveDocument.SaveAs(new_file_abs, FileFormat=16)
            doc.Close()
            word.Quit()

            if os.path.exists(doc_path):
                os.remove(doc_path)

            return new_file_abs
        except Exception as e:
            logging.error(f"Error converting {doc_path} to .docx: {e}")
            return None

    def separate_natural_language_from_code(self, text):
        """
        将输入的文本划分为自然语言文本和代码，返回自然语言文本部分和代码部分。
        参数: text (str): 输入的全部文本。
        返回: tuple (str, str): 包含自然语言和代码的文本部分。
        """
        natural_language_text = []
        code_text = []
        lines = text.split('\n')
        for line in lines:
            if re.search("[\u4e00-\u9fff]", line):  # 检测是否包含中文字符
                natural_language_text.append(line)
            else:
                code_text.append(line)  # 将代码行加入code_text
                logging.info(f"Identified as code or irrelevant: {line}")

        # 返回自然语言文本和代码文本
        return '\n'.join(natural_language_text), '\n'.join(code_text)

    def add_document(self, file_path):
        if file_path.endswith('.doc'):
            file_path = self.doc_to_docx(file_path)

        document_text = self.read_docx(file_path)
        filename = os.path.splitext(os.path.basename(file_path))[0]
        if document_text is not None:
            document_text = self.preprocess_text(document_text)
            # 分离自然语言和代码
            natural_language_text, code_text = self.separate_natural_language_from_code(document_text)

            self.documents[filename] = {
                '姓名': filename,
                '自然语言内容': natural_language_text,
                '代码内容': code_text,  # 存储代码部分
                '评分': None
            }
            logging.info(f"Added document with separated content: {filename}")
        else:
            logging.error(f"Failed to process file: {file_path}")
            self.documents[filename] = {
                '姓名': filename,
                '自然语言内容': 'error',
                '代码内容': 'error',
                '评分': None
            }

    def process_folder(self, folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            self.add_document(file_path)
        # 在此处调用相似度计算函数，更新self.documents中每个文档的'similarity'键



